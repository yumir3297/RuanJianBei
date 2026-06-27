from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "交付物-产品设计文档大纲.md"
OUTPUT = ROOT / "docs" / "交付物-产品设计文档.docx"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "17324D"
MUTED = "6B7280"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
BLACK = "111827"
GOLD = "9A6A22"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color="CBD5E1", size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def choose_widths(rows: list[list[str]]) -> list[int]:
    cols = len(rows[0])
    if cols == 1:
        return [CONTENT_WIDTH_DXA]
    if cols == 2:
        first_max = max(len(row[0]) for row in rows)
        second_max = max(len(row[1]) for row in rows)
        ratio = first_max / max(first_max + second_max, 1)
        first = int(CONTENT_WIDTH_DXA * min(max(ratio, 0.22), 0.38))
        return [first, CONTENT_WIDTH_DXA - first]

    lengths = []
    for col in range(cols):
        longest = max(len(row[col]) for row in rows)
        average = sum(len(row[col]) for row in rows) / len(rows)
        lengths.append(max(longest * 0.35 + average * 0.65, 6))

    minimums = [900] * cols
    if cols == 3:
        minimums = [1200, 2300, 2300]
    elif cols == 4:
        minimums = [900, 1900, 1200, 2200]

    available = CONTENT_WIDTH_DXA - sum(minimums)
    weights = [max(length - 6, 1) for length in lengths]
    total = sum(weights)
    widths = [
        minimums[idx] + int(available * weights[idx] / total)
        for idx in range(cols)
    ]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def set_run_font(
    run,
    *,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    east_asia: str = "Microsoft YaHei",
    latin: str = "Calibri",
) -> None:
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_pr.extend([r_fonts, color, underline])
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\([^)]+\))")


def add_inline(paragraph, text: str, *, size: float | None = None, color: str = BLACK) -> None:
    for token in INLINE_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(
                run,
                size=(size - 0.5 if size else 9.5),
                color=DARK_BLUE,
                east_asia="Microsoft YaHei",
                latin="Consolas",
            )
            run.font.highlight_color = None
        elif token.startswith("[") and "](" in token and token.endswith(")"):
            label, url = token[1:].split("](", 1)
            add_hyperlink(paragraph, label, url[:-1])
        else:
            run = paragraph.add_run(token)
            set_run_font(run, size=size, color=color)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_specs = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    styles["Heading 1"].paragraph_format.page_break_before = True


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    set_run_font(run, size=9, color=MUTED)
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, fld_sep, fld_end])


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("灵山胜境 AI 数字人导览系统  |  产品设计文档")
    set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    add_page_number(p)


def add_cover(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("第十四届中国软件杯 A5 赛题参赛作品")
    set_run_font(run, size=11, color=GOLD, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("灵山胜境 AI 数字人导览系统")
    set_run_font(run, size=28, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(64)
    run = subtitle.add_run("产 品 设 计 文 档")
    set_run_font(run, size=18, color=BLUE, bold=True)

    meta = [
        ("文档版本", "V0.9（初赛设计稿）"),
        ("编制日期", "2026 年 6 月"),
        ("学校 / 团队", "待补充"),
        ("文档状态", "主体已完成，截图、团队信息和最终验收数据待补"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [1800, 5000])
    set_table_borders(table, color="D7DEE8", size="4")
    for idx, (label, value) in enumerate(meta):
        left, right = table.rows[idx].cells
        for cell in (left, right):
            set_cell_margins(cell, top=110, bottom=110, start=140, end=140)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(left, LIGHT_BLUE)
        p = left.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(label)
        set_run_font(run, size=10.5, color=DARK_BLUE, bold=True)
        p = right.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(value)
        set_run_font(run, size=10.5, color=BLACK)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_contents(doc: Document, lines: list[str]) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(14)
    run = title.add_run("目录")
    set_run_font(run, size=20, color=NAVY, bold=True)

    entries = ["阅读说明"]
    for line in lines:
        if line.startswith("# ") and line != "# 灵山胜境 AI 数字人导览系统":
            entries.append(line[2:].strip())

    for entry in entries:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_together = True
        run = p.add_run(entry)
        set_run_font(run, size=11, color=DARK_BLUE, bold=entry.startswith("第 "))

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(14)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run("提示：在 Word 中可基于现有标题层级插入自动目录并更新页码。")
    set_run_font(run, size=9, color=MUTED, italic=True)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_numbering(doc: Document, *, ordered: bool, start: int = 1) -> int:
    numbering = doc.part.numbering_part.element
    existing_abs = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    existing_num = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(existing_abs, default=-1) + 1
    num_id = max(existing_num, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start_node = OxmlElement("w:start")
    start_node.set(qn("w:val"), str(start))
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1." if ordered else "•")
    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "100")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])

    level.extend([start_node, num_fmt, level_text, level_jc, p_pr])
    abstract.append(level)
    first_num_index = next(
        (
            index
            for index, child in enumerate(numbering)
            if child.tag == qn("w:num")
        ),
        len(numbering),
    )
    numbering.insert(first_num_index, abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), str(start))
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color="B9C8D8", size="5")
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    add_inline(p, text, size=10.5, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color="D8DEE8", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=90, bottom=90, start=140, end=140)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 0.95
    for idx, line in enumerate(lines):
        if idx:
            p.add_run("\n")
        run = p.add_run(line)
        set_run_font(
            run,
            size=8.0,
            color="263445",
            east_asia="Microsoft YaHei",
            latin="Consolas",
        )
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        row = [cell.strip() for cell in lines[idx].strip().strip("|").split("|")]
        rows.append(row)
        idx += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in rows[1]):
        del rows[1]
    return rows, idx


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    normalized = [row + [""] * (col_count - len(row)) for row in rows]
    table = doc.add_table(rows=len(normalized), cols=col_count)
    set_table_geometry(table, choose_widths(normalized))
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])

    body_size = 8.7 if col_count >= 4 else 9.2
    for r_idx, row in enumerate(normalized):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, "FAFBFC")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if r_idx == 0:
                p.paragraph_format.keep_with_next = True
            if col_count >= 3 and c_idx in (0, col_count - 1) and len(value) < 20:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(
                p,
                value,
                size=body_size,
                color=DARK_BLUE if r_idx == 0 else BLACK,
            )
            for run in p.runs:
                if r_idx == 0:
                    run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def build_document() -> None:
    source_lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    add_cover(doc)
    add_contents(doc, source_lines)

    start = 0
    for idx, line in enumerate(source_lines):
        if line.strip() == "## 阅读说明":
            start = idx
            break

    idx = start
    in_code = False
    code_lines: list[str] = []
    paragraph_buffer: list[str] = []
    bullet_num_id = add_numbering(doc, ordered=False)
    current_ordered_num_id: int | None = None
    previous_was_ordered = False
    first_content_heading = True

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        text = " ".join(item.strip() for item in paragraph_buffer).strip()
        if text:
            p = doc.add_paragraph()
            p.paragraph_format.widow_control = True
            add_inline(p, text, size=11)
        paragraph_buffer = []

    while idx < len(source_lines):
        raw = source_lines[idx]
        stripped = raw.strip()

        if stripped.startswith("```"):
            flush_paragraph()
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            idx += 1
            continue
        if in_code:
            code_lines.append(raw)
            idx += 1
            continue

        if not stripped:
            flush_paragraph()
            previous_was_ordered = False
            current_ordered_num_id = None
            idx += 1
            continue

        if stripped == "---":
            flush_paragraph()
            idx += 1
            continue

        if stripped.startswith("|"):
            flush_paragraph()
            rows, idx = parse_table(source_lines, idx)
            add_table(doc, rows)
            previous_was_ordered = False
            current_ordered_num_id = None
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading_match:
            flush_paragraph()
            hashes, text = heading_match.groups()
            if text in ("灵山胜境 AI 数字人导览系统", "产品设计文档"):
                idx += 1
                continue

            level = len(hashes)
            if text == "阅读说明":
                level = 1
            elif text.startswith("附录 "):
                level = 1

            p = doc.add_paragraph(style=f"Heading {level}")
            if first_content_heading:
                p.paragraph_format.page_break_before = False
                first_content_heading = False
            add_inline(p, text, size={1: 16, 2: 13, 3: 12}[level], color=BLUE if level < 3 else DARK_BLUE)
            for run in p.runs:
                run.bold = True
            previous_was_ordered = False
            current_ordered_num_id = None
            idx += 1
            continue

        if stripped.startswith("> "):
            flush_paragraph()
            add_callout(doc, stripped[2:].strip())
            idx += 1
            continue

        bullet_match = re.match(r"^-\s+(.+)$", stripped)
        if bullet_match:
            flush_paragraph()
            p = doc.add_paragraph()
            apply_num(p, bullet_num_id)
            add_inline(p, bullet_match.group(1), size=11)
            previous_was_ordered = False
            current_ordered_num_id = None
            idx += 1
            continue

        ordered_match = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if ordered_match:
            flush_paragraph()
            number = int(ordered_match.group(1))
            if not previous_was_ordered or current_ordered_num_id is None:
                current_ordered_num_id = add_numbering(doc, ordered=True, start=number)
            p = doc.add_paragraph()
            apply_num(p, current_ordered_num_id)
            add_inline(p, ordered_match.group(2), size=11)
            previous_was_ordered = True
            idx += 1
            continue

        paragraph_buffer.append(stripped)
        idx += 1

    flush_paragraph()

    core = doc.core_properties
    core.title = "灵山胜境 AI 数字人导览系统 - 产品设计文档"
    core.subject = "第十四届中国软件杯 A5 赛题参赛作品"
    core.author = "参赛团队"
    core.keywords = "中国软件杯, AI数字人, 景区导览, RAG, 多模态"
    core.comments = "由项目 Markdown 设计文档生成"

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
