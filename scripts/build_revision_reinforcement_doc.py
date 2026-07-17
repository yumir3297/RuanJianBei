from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "A5作品整改与补强建议书-2026-07-15.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "5B6472"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
LIGHT_RED = "FDECEC"
LIGHT_GOLD = "FFF7E6"
GREEN = "247A4D"
RED = "A52A2A"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = tc_pr.first_child_found_in("w:tcMar")
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for tag, value in (("top", "80"), ("bottom", "80"), ("start", "120"), ("end", "120")):
                node = margins.find(qn(f"w:{tag}"))
                if node is None:
                    node = OxmlElement(f"w:{tag}")
                    margins.append(node)
                node.set(qn("w:w"), value)
                node.set(qn("w:type"), "dxa")


def set_run(run, size=11, color=INK, bold=False, italic=False) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def format_paragraph(paragraph, before=0, after=6, line=1.1, align=None) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_text(doc, text: str, *, size=11, color=INK, bold=False, italic=False, before=0, after=6, line=1.1, align=None) -> None:
    p = doc.add_paragraph()
    format_paragraph(p, before, after, line, align)
    set_run(p.add_run(text), size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text: str, level: int) -> None:
    p = doc.add_paragraph()
    sizes = {1: 16, 2: 13, 3: 12}
    colors = {1: BLUE, 2: BLUE, 3: DARK_BLUE}
    before = {1: 16, 2: 12, 3: 8}[level]
    after = {1: 8, 2: 6, 3: 4}[level]
    format_paragraph(p, before, after, 1.1)
    p.style = f"Heading {level}"
    set_run(p.add_run(text), size=sizes[level], color=colors[level], bold=True)
    return p


def add_bullet(doc, text: str, *, level=0, color=INK) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.38 + level * 0.22)
    p.paragraph_format.first_line_indent = Inches(-0.19)
    format_paragraph(p, 0, 4, 1.2)
    set_run(p.add_run(text), size=10.5, color=color)
    return p


def add_number(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.38)
    p.paragraph_format.first_line_indent = Inches(-0.19)
    format_paragraph(p, 0, 4, 1.2)
    set_run(p.add_run(text), size=10.5)
    return p


def add_callout(doc, title: str, text: str, fill=LIGHT_GOLD, color="7A5A00") -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    format_paragraph(p, 0, 0, 1.15)
    set_run(p.add_run(f"{title}："), size=10.5, color=color, bold=True)
    set_run(p.add_run(text), size=10.5, color=INK)
    add_text(doc, "", size=2, after=2)


def add_matrix(doc, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header_cells = table.rows[0].cells
    for cell, text in zip(header_cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        format_paragraph(p, 0, 0, 1.05, WD_ALIGN_PARAGRAPH.CENTER)
        set_run(p.add_run(text), size=9.5, color=DARK_BLUE, bold=True)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, (cell, text) in enumerate(zip(cells, row_values)):
            p = cell.paragraphs[0]
            format_paragraph(p, 0, 0, 1.08)
            color = RED if text == "P0" else INK
            bold = text in {"P0", "P1", "P2", "P3"}
            set_run(p.add_run(text), size=9.2, color=color, bold=bold)
            if idx == 0 and text.startswith("P0"):
                set_cell_shading(cell, LIGHT_RED)
    add_text(doc, "", size=2, after=2)


def add_location_block(doc, locations: list[str]) -> None:
    p = doc.add_paragraph()
    format_paragraph(p, 0, 5, 1.15)
    set_run(p.add_run("问题位置："), size=10.5, color=DARK_BLUE, bold=True)
    for index, location in enumerate(locations):
        if index:
            set_run(p.add_run("；"), size=10.5, color=MUTED)
        set_run(p.add_run(location), size=9.5, color=MUTED)


def add_issue(doc, number: int, priority: str, title: str, locations: list[str], issue: str, score: str, solution: list[str], acceptance: list[str], demo: str) -> None:
    add_heading(doc, f"{number}. {priority} - {title}", 2)
    add_location_block(doc, locations)
    add_text(doc, issue, size=10.5, after=4)
    add_text(doc, f"扣分逻辑：{score}", size=10.5, color=RED if priority == "P0" else INK, bold=True, after=4)
    add_text(doc, "建议修改方案", size=10.5, color=DARK_BLUE, bold=True, before=2, after=2)
    for item in solution:
        add_bullet(doc, item)
    add_text(doc, "验收标准", size=10.5, color=DARK_BLUE, bold=True, before=2, after=2)
    for item in acceptance:
        add_bullet(doc, item, color=GREEN)
    add_text(doc, f"答辩呈现：{demo}", size=10, color=MUTED, italic=True, before=1, after=6)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        styles[name].font.name = "Microsoft YaHei"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    format_paragraph(header, 0, 0, 1.0)
    set_run(header.add_run("中国软件杯 A5 作品整改与补强建议书"), size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_paragraph(footer, 0, 0, 1.0)
    set_run(footer.add_run("内部整改依据 | 2026-07-15"), size=8.5, color=MUTED)


def build() -> None:
    doc = Document()
    configure_document(doc)

    add_text(doc, "中国软件杯 A5 作品", size=12, color=MUTED, bold=True, after=2)
    add_text(doc, "整改与补强建议书", size=25, color=INK, bold=True, after=5)
    add_text(doc, "适用对象：灵山胜境景区 AI 数字人导览系统 | 用途：提交前整改、任务拆分、答辩验收", size=11.5, color=MUTED, after=14)
    metadata = doc.add_table(rows=4, cols=2)
    metadata.style = "Table Grid"
    set_table_geometry(metadata, [2700, 6660])
    meta_rows = [
        ("文档定位", "以当前代码、构建、接口烟测和回归测试为依据的整改作战手册"),
        ("评审基线", "前端生产构建通过；后端 93 通过、2 失败；路线后端接口可用；语音与动态路线未完成端到端验收"),
        ("使用原则", "没有可复验结果的能力，一律不作为完成项写入答辩材料"),
        ("整改目标", "先补赛题硬要求与业务闭环，再提高现场可演示性和作品差异化"),
    ]
    for row, (label, value) in zip(metadata.rows, meta_rows):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        for cell, text, bold in ((row.cells[0], label, True), (row.cells[1], value, False)):
            p = cell.paragraphs[0]
            format_paragraph(p, 0, 0, 1.1)
            set_run(p.add_run(text), size=10, color=DARK_BLUE if bold else INK, bold=bold)
    add_text(doc, "", size=2, after=3)
    add_callout(doc, "本次核心判断", "系统的产品外观和模块数量已具备竞争力，但高分取决于三条真正打通的证据链：真实路线推荐、真实语音问答、可复验准确率与稳定性。", fill=LIGHT_GOLD)

    add_heading(doc, "一、整改总策略", 1)
    add_text(doc, "不要再平均投入所有页面。应按“赛题硬要求 -> 后端闭环 -> 可视化演示 -> 体验打磨”的顺序推进。任何会被评委一句“这是真的吗？”击穿的功能，都必须优先补证据。", size=10.8)
    add_matrix(
        doc,
        ["优先级", "目标", "完成定义", "建议投入"],
        [
            ["P0", "消除硬伤", "评委从页面操作可得到真实接口结果；测试全绿", "提交前立即完成"],
            ["P1", "补齐硬指标", "语音、动态路线、准确率均有可复验报告", "提交前必须完成"],
            ["P2", "形成差异化", "反馈运营闭环和可信数据表达可现场展示", "完成后明显加分"],
            ["P3", "提升质感", "安全、性能与视觉细节优化", "时间充足再做"],
        ],
        [1100, 2100, 3900, 2260],
    )
    add_text(doc, "推荐的最小冲刺路径", size=11, color=DARK_BLUE, bold=True, before=4, after=3)
    for step in [
        "先修路线页和回归测试，确保“页面操作 -> 后端计算 -> 数据库记录 -> 页面展示”没有断点。",
        "再用固定测试集复跑问答评测，并保留原始结果、失败样例与版本号。",
        "最后做语音和 Coze 动态路线的双路径烟测：成功路径与降级路径都要录屏。",
    ]:
        add_number(doc, step)

    add_heading(doc, "二、P0 必须立即解决的问题", 1)
    add_issue(
        doc, 1, "P0", "路线主页面仍为静态预设，未接入真实推荐接口",
        ["frontend/src/views/tourist/RouteView.vue:149", "frontend/src/views/tourist/RouteView.vue:290", "frontend/src/api/recommend.js:3", "backend/app/api/recommend.py:14"],
        "后端 /api/recommend/ 已能依据会话、兴趣、同行人和时长返回数据库路线，但路线主页面仍直接过滤 PRESET_ROUTES。评委在最显眼的“路线规划”入口看到的仍是固定内容。",
        "这会把“个性化路线推荐”降级为静态展示，直接影响赛题硬要求与实际业务闭环得分。",
        [
            "在 RouteView 中导入 fetchRecommendations，并新增 sessionId、兴趣、时长、同行人、避拥挤偏好状态。",
            "用户提交偏好后调用真实接口，用返回的 routes 替换 PRESET_ROUTES；接口异常时才展示明确标注的本地备选方案。",
            "将 route_id、reason、guide_points、experiences、source 全部渲染到详情页；不要只渲染标题和固定文案。",
            "将本次推荐与 session_id 绑定，确认后端 VisitorProfile 已记录本次兴趣偏好。",
        ],
        [
            "修改前后抓包可见页面真实发出 POST /api/recommend/，请求包含 session_id。",
            "同一页面输入“亲子 4 小时”与“文化深度 6 小时”后，返回路线标题、理由、停靠点至少有一项不同。",
            "断开后端后页面不伪装为个性化结果，而是显示“本地备选路线/服务暂不可用”。",
            "推荐成功后数据库 visitor_profiles 中存在该 session 的兴趣或人群标签。",
        ],
        "现场输入两组偏好，展示路线即时变化，再打开后台或日志说明偏好已沉淀。",
    )
    add_issue(
        doc, 2, "P0", "后端全量回归测试未通过",
        ["backend/app/tests/test_blind_spot_admin.py:61", "backend/app/tests/test_blind_spot_admin.py:142", "backend/app/api/admin.py:51", "backend/app/api/auth.py:56"],
        "当前后台接口已增加真实 Bearer 鉴权，但盲区管理测试仍以未认证方式访问，导致 pytest 为 93 passed、2 failed。功能方向正确，测试契约没有同步。",
        "评委或指导老师一旦运行测试，失败结果会削弱系统稳定性和工程规范的可信度。",
        [
            "在测试夹具中先请求 /api/auth/login，再为 admin 接口附加 Authorization: Bearer <token>。",
            "分别保留未认证返回 401、已认证成功、无效令牌返回 401 三类测试，不能只为了通过测试而移除鉴权。",
            "在 CI 或提交前检查中固定执行 backend 全量 pytest。",
        ],
        [
            "backend/.venv/Scripts/python.exe -m pytest app/tests -q 输出 0 failed。",
            "盲区列表、盲区转 FAQ、知识管理和头像管理均有已认证成功与未认证拒绝测试。",
            "失败时输出可定位的错误，不依赖人工临时修改数据库。",
        ],
        "答辩备用页展示测试汇总，不需要展示全部日志，但必须保证现场重跑为全绿。",
    )
    add_issue(
        doc, 3, "P0", "景区事实问答 90% 指标没有当前可复验结论",
        ["docs/100题评测修复-任务书-Claude.md:10", "docs/硬指标达标确认.md:14", "eval/testset/e2e_qa_100.json", "eval/scripts/run_e2e_chat_eval.py"],
        "仓库同时存在“100 题准确率 64%”和“10 题抽样 90%”的描述，后者不能替代赛题要求的完整、当前版本评测。预测达到 85-90% 也不能算验收。",
        "准确率是赛题硬指标。证据冲突时，严格评委会按最低且可复验的结果判断。",
        [
            "冻结当前数据包、FAQ、RAG 索引、模型版本和环境配置，给评测报告写明版本号与执行时间。",
            "运行 100 题端到端评测，逐题保存 query、命中链路、来源、答案、耗时、判定与失败原因。",
            "将“关键词未命中但语义正确”和“模型调用失败”拆分统计，不能用解释性文字掩盖总指标。",
            "如未达到 90%，先补齐高频事实 FAQ 和知识块，再复跑；答辩时只展示最终可复现版本。",
        ],
        [
            "报告总题数为 100，准确率 >= 90%，并可导出失败题清单。",
            "同一测试集连续执行两次，准确率波动不超过 3 个百分点；网络失败须单独计入稳定性，不得静默忽略。",
            "每题均能看到至少一个官方资料来源或明确拒答原因。",
        ],
        "展示一页“100题评测卡”：准确率、证据率、P50/P95、版本号和 3 个典型样例。",
    )

    add_heading(doc, "三、P1 提交前必须解决的问题", 1)
    add_issue(
        doc, 4, "P1", "动态路线 Coze 分支未完成端到端联调",
        ["backend/app/services/qa/pipeline.py", "backend/app/services/coze/client.py", "docs/2026-07-15-dynamic-route-planning-integration-for-deepseek-review.md"],
        "后端已有意图识别、Coze 调用、白名单校验和本地降级，但项目说明明确写明仍需真实前端请求完成端到端联调；动态结果也尚未以结构化路线卡片呈现。",
        "有代码不等于现场可用。若 Coze 超时、令牌无效或返回不合规数据，评委会关注是否真正降级而不是页面卡死。",
        [
            "新增动态路线烟测脚本：成功调用、错误 Token、超时、未授权景点 ID、空 route_stops 五种场景。",
            "将 route_stops、adjustments、warning 通过独立 SSE 事件返回，前端渲染站点卡、原因和实时数据来源。",
            "强制要求 route_stops 非空，每个停靠点都必须含 attraction_id 与 reason；校验失败直接降级。",
            "实时天气/客流若仍是 mock，必须在结果区域和答辩口播中醒目标明“演示模拟数据”。",
        ],
        [
            "真实浏览器输入“带父母两小时、避开拥挤”后，日志出现 dynamic_route_coze 或明确的降级原因。",
            "人为关闭 Coze 后，页面在可接受时间内回到本地路线/问答，不出现空白、报错栈或无限加载。",
            "Coze 返回越权景点时，前端不会展示该景点。",
        ],
        "连续演示成功路线和断网降级路线，证明不是单点依赖。",
    )
    add_issue(
        doc, 5, "P1", "语音问答缺少真实音频、时延和稳定性验收",
        ["backend/.env（ASR_PROVIDER 重复）", "backend/app/api/voice.py:17", "frontend/src/views/tourist/ChatView.vue", "frontend/src/composables/useRecorder.js"],
        "当前运行时 ASR 为 QwenASRService，且 ASR、LLM、TTS 凭据均已配置；但没有基于真实音频的转写准确率、端到端时延和失败降级证据。配置文件同时有重复 ASR_PROVIDER。",
        "赛题要求语音问答小于 5 秒。浏览器 Web Speech 或占位回退不能替代后端真实服务验证。",
        [
            "清理 .env，仅保留一个 ASR_PROVIDER=qwen，并在启动日志打印 provider 名称而不打印密钥。",
            "准备至少 5 段真实录音，覆盖普通话、景区专名、短问题、长问题和噪声场景。",
            "记录录音结束、ASR 返回、首字显示、首段语音播放、完整回答结束五个时间点。",
            "为 ASR、LLM、TTS 分别设置超时与明确降级提示；回退浏览器语音时应显示“浏览器播报”，不可伪装为云端 TTS。",
        ],
        [
            "5 段语音均能转写为可编辑文本；景区核心专名的转写正确率达到预设目标。",
            "FAQ 和 RAG 两类问题分别统计端到端 P50/P95，P95 <= 5 秒或在答辩中如实说明未达到。",
            "关闭任一外部服务后，页面能提示降级而不崩溃。",
        ],
        "用一段现场录音完整演示“说话 -> 转写 -> 引用回答 -> 数字人播报”，并展示一次明确降级。",
    )
    add_issue(
        doc, 6, "P1", "游客反馈未关联具体问答记录，运营闭环不够精确",
        ["frontend/src/stores/chat.js:submitFeedback", "frontend/src/views/tourist/ChatView.vue:55", "backend/app/api/insights.py:180", "backend/app/models/visitor_feedback.py:21"],
        "前端已从 SSE done 事件得到 message.chatLogId，但 submitFeedback 目前只提交 session_id 与 rating，未传 chat_log_id。后端因此无法将“需改进”精确定位到某条答案、来源和命中链路。",
        "游客体验报告虽已真实聚合，但不能把负面反馈变成知识修复任务，会被视为数据展示而非运营闭环。",
        [
            "提交反馈时传入 chat_log_id: message.chatLogId，并在未拿到 ID 时禁用反馈或提示稍后再试。",
            "后台报告按问题、命中层级、来源文档、负面原因聚合，点击问题可跳转到聊天日志。",
            "为低评分或负面反馈生成待处理任务，管理员可转 FAQ、补知识、标记已解决。",
        ],
        [
            "每条反馈数据库均含 chat_log_id，且能反查同一 session 的原问题和回答。",
            "后台可筛选“回答不准确”的具体问答，而非只显示总数。",
            "一条负反馈经补 FAQ 后重新提问，命中层级从 rag/盲区变为 FAQ 或有明确修复记录。",
        ],
        "现场演示“点踩 -> 后台定位 -> 转 FAQ -> 再次提问命中”的一分钟闭环。",
    )

    add_heading(doc, "四、P2/P3 明显加分但不应抢占主线的问题", 1)
    add_issue(
        doc, 7, "P2", "体验报告样本量过小，满意度容易被误读",
        ["backend/app/services/experience_report.py:26", "frontend/src/views/admin/ExperienceReport.vue:44", "backend/app/db/bootstrap.py:188"],
        "当前体验报告已区分 real/demo，并展示反馈覆盖率；但最近数据约 125 次交互仅 1 条反馈，100% 满意度没有统计显著性。",
        "如果在答辩中把低覆盖率满意度包装为运营成果，评委会质疑数据真实性。",
        [
            "保留 data_mode 与反馈覆盖率，满意度旁增加“样本不足”阈值提示，例如反馈 < 20 条时不显示结论性趋势。",
            "准备 15-30 条可追溯的演示反馈，标记为评委演示样本，不冒充真实景区运营数据。",
            "将体验报告的核心价值改为“问题发现和改进建议”，而非夸大满意度。",
        ],
        [
            "数据模式、样本数、覆盖率在页面第一屏可见。",
            "低于阈值时展示“样本不足，仅供趋势观察”。",
            "演示数据与真实数据在接口、页面和讲稿中均可区分。",
        ],
        "用报告中的高频问题和负向原因引出知识盲区闭环，而不是强调单一满意度百分比。",
    )
    add_issue(
        doc, 8, "P2", "后台认证已真实接入，但令牌安全策略仍较弱",
        ["backend/app/api/auth.py:28", "backend/app/api/auth.py:39", "backend/app/core/config.py:91"],
        "后台从前端硬编码密码升级为后端 HMAC 令牌，这是有效改进；但令牌校验未看到过期时间，默认密码和默认密钥仍在代码默认配置中。",
        "这不一定影响竞赛演示，但会影响工程规范、稳定性和安全性评价。",
        [
            "令牌 payload 加入 expires_at 并校验；设置例如 8 小时有效期。",
            "生产/演示环境要求 ADMIN_PASSWORD 与 ADMIN_TOKEN_SECRET 必须由 .env 提供，未设置时拒绝启动或输出警告。",
            "补登录成功、错误密码、令牌过期、退出后失效的自动化测试。",
        ],
        [
            "过期令牌访问 /api/admin/* 返回 401。",
            "默认密钥不再用于演示环境，启动自检可以发现错误配置。",
            "现有后台管理功能在鉴权收紧后仍全部可用。",
        ],
        "不必作为主讲内容，但可在被问到权限管理时展示真实后端鉴权。",
    )
    add_issue(
        doc, 9, "P3", "前端生产包偏大，首次加载可能影响现场体验",
        ["frontend npm run build 输出：index 约 1.0 MB，avatarConfig 约 793 KB", "frontend/src/router/index.js"],
        "生产构建已通过，但 Vite 仍提示多个压缩前 chunk 超过 500 KB。3D 数字人和可视化依赖可能增加首次加载等待。",
        "不是核心扣分项，但现场网络不稳定时会放大首屏白屏和加载卡顿风险。",
        [
            "保持路由懒加载；将数字人模型、ECharts、管理后台配置页拆为动态 import。",
            "为数字人区域提供加载骨架和超时降级，不阻塞文字问答。",
            "在答辩电脑上预热一次资源，并准备本地部署或可靠网络。",
        ],
        [
            "首次进入游客问答页可先发送文字问题，数字人资源未加载不阻塞交互。",
            "构建告警显著减少或明确记录为可接受的 3D 资源成本。",
            "弱网下页面有加载反馈，不出现长时间空白。",
        ],
        "现场先进入首页预热资源，再开始正式演示；不要让评委等待大模型与 3D 模型同时加载。",
    )

    add_heading(doc, "五、最值得强化的三条作品亮点", 1)
    add_text(doc, "这三项应被组织成一条连续故事，而不是分散展示页面：游客得到可信讲解，系统理解真实偏好，运营人员据此持续改进。", size=10.8)
    add_matrix(
        doc,
        ["亮点", "为什么有竞争力", "必须补强的证据", "推荐答辩表达"],
        [
            ["证据驱动的景区讲解", "回答附来源、知识库与盲区管理形成可信 AI，而不只是通用聊天", "100题评测、引用可点开、错误能进入盲区", "系统不靠模型自由编造，每次讲解都有景区资料依据。"],
            ["数字人多模态导览", "语音、播报、口型、表情与图片识别构成直观可见的体验差异", "真实 ASR/TTS 时延和失败降级录屏", "游客可以说、可以问、可以看，数字人能以可解释的内容回应。"],
            ["运营反馈闭环", "从游客点踩到知识修复可体现系统可持续运营", "反馈绑定 chat_log、转 FAQ、复问验证", "我们不仅回答问题，还能持续发现并修复服务盲区。"],
        ],
        [1500, 2800, 2750, 2310],
    )

    add_heading(doc, "六、提交前验收清单", 1)
    add_text(doc, "以下项目必须由实际运行结果签字确认。未通过的项目不可在 PPT 或视频中写成“已完成”。", size=10.8)
    checklist = [
        "前端：npm run build 成功，路线页发出真实推荐请求并根据返回结果变化。",
        "后端：pytest app/tests -q 为全绿，后台未认证访问返回 401、认证访问成功。",
        "问答：100 题评测准确率 >= 90%，保留失败题和版本化报告。",
        "语音：5 段真实录音完成 ASR -> 问答 -> TTS，保留 P50/P95 时延。",
        "动态路线：Coze 成功、超时、错误 Token、非法景点、关闭服务五种场景均有结果。",
        "反馈：任一“需改进”可定位到 chat_log，并能进入知识盲区/FAQ 修复闭环。",
        "数据：演示数据、模拟实时数据、真实反馈数据均在页面和讲稿中明确标识。",
        "演示：按 7 分钟流程完整走两遍，第二遍计时并在无网络或网络抖动条件下验证降级。",
    ]
    for item in checklist:
        add_bullet(doc, item, color=GREEN)

    add_callout(doc, "最终提交门槛", "只要路线主页面、测试全绿、100题评测、语音烟测四项没有完成，就不要把作品定位为“全部满足硬性要求”。这些是从 70 分段进入高分段的关键分水岭。", fill=LIGHT_RED, color=RED)

    add_heading(doc, "七、建议的 7 分钟演示结构", 1)
    schedule = [
        ["0:00-0:40", "问题与场景", "说明景区讲解、路线决策和运营反馈三类痛点。"],
        ["0:40-2:20", "可信问答", "语音或文字问灵山景点，展示来源引用、数字人播报与图片识别。"],
        ["2:20-3:50", "个性路线", "输入不同人群/时长偏好，展示真实推荐结果变化和动态路线降级。"],
        ["3:50-5:10", "后台运营", "展示知识库、数字人配置、反馈报告和盲区列表。"],
        ["5:10-6:10", "闭环修复", "点踩一条回答，后台定位，转 FAQ 或补知识，再次提问验证。"],
        ["6:10-7:00", "硬指标与总结", "展示测试全绿、100题评测、语音时延，并总结作品特色。"],
    ]
    add_matrix(doc, ["时间", "演示段", "必须证明的事实"], schedule, [1350, 2200, 5810])

    add_heading(doc, "附录：本建议书的证据边界", 1)
    add_text(doc, "本建议书基于 2026-07-15 工作区代码、构建和接口烟测形成。已验证的内容包括前端生产构建、部分 FastAPI 接口、Qwen ASR 实例和后端路线接口。未在本次复核中完成外部模型网络调用、Coze 浏览器端到端演示、100题当前版本复跑和演示视频核验，因此这些项目被列为待验收，而非默认完成。", size=10.5)
    add_text(doc, "主要证据文件：frontend/src/views/tourist/RouteView.vue；frontend/src/stores/chat.js；backend/app/api/auth.py；backend/app/api/voice.py；backend/app/services/qa/pipeline.py；backend/app/services/experience_report.py；docs/100题评测修复-任务书-Claude.md；docs/2026-07-15-dynamic-route-planning-integration-for-deepseek-review.md。", size=9.2, color=MUTED, after=0)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
