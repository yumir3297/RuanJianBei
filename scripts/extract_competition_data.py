from __future__ import annotations

import json
import re
from pathlib import Path
from statistics import mean

import pandas as pd
from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parents[1]
RAW_DIR = PROJECT_ROOT / "data" / "raw"
PROCESSED_DIR = PROJECT_ROOT / "data" / "processed"


STRUCTURED_FIELDS = [
    "景区名称",
    "景点ID",
    "景点名称",
    "具体位置",
    "建筑/景观参数",
    "核心功能",
    "文化内涵",
    "详细介绍",
    "游玩亮点",
    "演艺/开放信息",
    "备注",
]

ROUTE_HEADING_RE = re.compile(r"^(?P<title>.+路线)（(?P<duration>.+)）$")
GENERIC_ROUTE_SECTION_TITLES = {"讲解重点", "特色体验"}
TIMELINE_MENTION_RE = re.compile(r"^\d{4}年")


def clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\n", " ").strip())


def sentence_join(parts: list[str]) -> str:
    filtered = [clean_text(part) for part in parts if clean_text(part)]
    return " ".join(filtered)


def truncate(text: str, max_length: int = 260) -> str:
    normalized = clean_text(text)
    if len(normalized) <= max_length:
        return normalized
    return normalized[: max_length - 1].rstrip() + "…"


def load_docx_by_keyword(keyword: str) -> Path:
    for path in RAW_DIR.glob("*.docx"):
        if keyword in path.name:
            return path
    raise FileNotFoundError(f"Cannot find docx file with keyword: {keyword}")


def load_xlsx() -> Path:
    files = list(RAW_DIR.glob("*.xlsx"))
    if not files:
        raise FileNotFoundError("Cannot find xlsx source file in data/raw.")
    return files[0]


def load_clean_paragraphs(docx_path: Path) -> list[str]:
    doc = Document(str(docx_path))
    return [clean_text(paragraph.text) for paragraph in doc.paragraphs if clean_text(paragraph.text)]


def extract_spot_records(structured_docx: Path) -> list[dict]:
    doc = Document(str(structured_docx))
    records: list[dict] = []

    for table in doc.tables:
        headers = [clean_text(cell.text) for cell in table.rows[0].cells]
        if headers != STRUCTURED_FIELDS:
            continue

        for row in table.rows[1:]:
            values = [clean_text(cell.text) for cell in row.cells]
            if not any(values):
                continue
            records.append(dict(zip(headers, values)))

    return records


def is_heading(text: str, next_text: str | None) -> bool:
    if not text:
        return False
    if ROUTE_HEADING_RE.match(text):
        return True
    if text.endswith("：") and len(text) <= 20:
        return True
    punctuation = "。；！？!?，,"
    if len(text) <= 30 and not any(symbol in text for symbol in punctuation):
        return bool(next_text and len(clean_text(next_text)) >= 20)
    if "：" in text and len(text) <= 28 and bool(next_text and len(clean_text(next_text)) >= 20):
        return True
    return False


def extract_guide_sections(guide_docx: Path) -> list[dict]:
    paragraphs = load_clean_paragraphs(guide_docx)
    if not paragraphs:
        return []

    sections: list[dict] = []
    current_heading = paragraphs[0]
    current_body: list[str] = []

    for index, text in enumerate(paragraphs[1:], start=1):
        next_text = paragraphs[index + 1] if index + 1 < len(paragraphs) else None
        if is_heading(text, next_text):
            if current_body:
                sections.append(
                    {
                        "title": current_heading.rstrip("："),
                        "content": sentence_join(current_body),
                    }
                )
            current_heading = text.rstrip("：")
            current_body = []
        else:
            current_body.append(text)

    if current_body:
        sections.append({"title": current_heading.rstrip("："), "content": sentence_join(current_body)})

    return sections


def extract_routes(guide_docx: Path) -> list[dict]:
    paragraphs = load_clean_paragraphs(guide_docx)
    routes: list[dict] = []
    stop_titles = {"实用游览贴士：全方位保障你的灵山之旅", "实用游览贴士"}

    idx = 0
    while idx < len(paragraphs):
        match = ROUTE_HEADING_RE.match(paragraphs[idx])
        if not match:
            idx += 1
            continue

        route = {
            "id": f"route_{len(routes) + 1:03d}",
            "title": match.group("title"),
            "duration_label": match.group("duration"),
            "route_plan": "",
            "guide_points": [],
            "experiences": [],
            "source": f"{guide_docx.name} - {paragraphs[idx]}",
        }
        mode: str | None = None
        idx += 1

        while idx < len(paragraphs):
            text = paragraphs[idx]
            if ROUTE_HEADING_RE.match(text) or text in stop_titles:
                break

            if text.startswith("路线规划："):
                route["route_plan"] = text.split("：", 1)[1].strip()
                mode = None
            elif text.startswith("讲解重点"):
                mode = "guide_points"
            elif text.startswith("特色体验"):
                mode = "experiences"
            else:
                if mode == "guide_points":
                    route["guide_points"].append(text)
                elif mode == "experiences":
                    route["experiences"].append(text)
            idx += 1

        routes.append(route)

    return routes


def score_spot_mention(spot_name: str, paragraph: str) -> int:
    score = 0

    if paragraph.startswith("路线规划："):
        return -100
    if paragraph.startswith(f"{spot_name}：") or paragraph.startswith(f"{spot_name}:"):
        score += 5
    elif paragraph.startswith(spot_name):
        score += 4

    if any(keyword in paragraph for keyword in ("重点讲解", "解析", "介绍", "亮点", "体验")):
        score += 2
    if TIMELINE_MENTION_RE.match(paragraph) or any(
        keyword in paragraph for keyword in ("工程奠基", "落成开光", "正式开放", "建成开放")
    ):
        score -= 2

    return score


def build_guide_mentions(paragraphs: list[str], spot_names: list[str]) -> dict[str, list[str]]:
    mentions: dict[str, list[tuple[int, str]]] = {spot_name: [] for spot_name in spot_names}

    for index, paragraph in enumerate(paragraphs):
        if len(paragraph) < 8:
            continue
        for spot_name in spot_names:
            if spot_name in paragraph:
                mentions[spot_name].append((index, paragraph))

    ranked_mentions: dict[str, list[str]] = {}
    for spot_name, candidates in mentions.items():
        ordered = sorted(
            candidates,
            key=lambda item: (-score_spot_mention(spot_name, item[1]), item[0]),
        )
        deduped: list[str] = []
        seen: set[str] = set()
        for _, paragraph in ordered:
            if paragraph in seen:
                continue
            seen.add(paragraph)
            deduped.append(paragraph)
        ranked_mentions[spot_name] = deduped

    return ranked_mentions


def spot_record_to_knowledge_entry(record: dict, source_name: str, guide_mentions: dict[str, list[str]]) -> dict:
    attraction_name = record["景点名称"]
    scenic_area = record["景区名称"]

    content_parts = []
    for field in STRUCTURED_FIELDS[3:]:
        value = record.get(field, "")
        if value:
            content_parts.append(f"{field}：{value}")

    supplement = sentence_join(guide_mentions.get(attraction_name, [])[:2])
    if supplement:
        content_parts.append(f"导览资料补充：{supplement}")

    aliases = [attraction_name]
    if scenic_area not in attraction_name:
        aliases.append(f"{scenic_area}{attraction_name}")

    return {
        "title": attraction_name,
        "category": "景点信息",
        "content": sentence_join(content_parts),
        "source": f"{source_name} - {scenic_area} / {attraction_name}",
        "aliases": aliases,
        "metadata": {
            "scenic_area": scenic_area,
            "attraction_id": record["景点ID"],
            "location": record["具体位置"],
            "parameters": record["建筑/景观参数"],
            "core_function": record["核心功能"],
            "cultural_meaning": record["文化内涵"],
            "opening_info": record["演艺/开放信息"],
            "remark": record["备注"],
        },
    }


def guide_section_to_knowledge_entry(section: dict, source_name: str) -> dict:
    title = section["title"]
    category = "历史文化"
    if "路线" in title:
        category = "游览路线"
    elif "贴士" in title or title in {"门票与优惠政策", "最佳游览时间", "餐饮", "住宿", "其他实用建议"}:
        category = "实用贴士"
    elif "景点" in title:
        category = "景点解读"

    return {
        "title": title,
        "category": category,
        "content": section["content"],
        "source": f"{source_name} - {title}",
        "aliases": [title],
        "metadata": {},
    }


def build_overview_faq(guide_sections: list[dict], source_name: str) -> list[dict]:
    faq_entries: list[dict] = []
    if not guide_sections:
        return faq_entries

    title_to_content = {section["title"]: section["content"] for section in guide_sections}
    overview = title_to_content.get("景区概况与千年历史渊源", "")
    if overview:
        faq_entries.append(
            {
                "id": "faq_overview_001",
                "category": "景区概况",
                "aliases": ["灵山胜境在哪里", "灵山胜境是什么景区", "灵山胜境介绍"],
                "answer": truncate(overview, 320),
                "sources": [f"{source_name} - 景区概况与千年历史渊源"],
            }
        )

    tips = title_to_content.get("最佳游览时间", "")
    if tips:
        faq_entries.append(
            {
                "id": "faq_tips_001",
                "category": "游览贴士",
                "aliases": ["灵山胜境最佳游览时间", "灵山胜境什么时候去最好"],
                "answer": truncate(tips, 260),
                "sources": [f"{source_name} - 最佳游览时间"],
            }
        )

    return faq_entries


def build_spot_faqs(
    spot_records: list[dict],
    source_name: str,
    guide_mentions: dict[str, list[str]],
) -> list[dict]:
    faq_entries: list[dict] = []
    counter = 1

    for record in spot_records:
        spot_name = record["景点名称"]
        scenic_area = record["景区名称"]
        base_source = f"{source_name} - {scenic_area} / {spot_name}"
        supplement = sentence_join(guide_mentions.get(spot_name, [])[:2])

        intro_parts = [
            f"{spot_name}位于{record['具体位置']}" if record["具体位置"] else "",
            record["详细介绍"],
            record["核心功能"],
            record["文化内涵"],
            record["游玩亮点"],
            supplement,
        ]
        intro_answer = truncate(sentence_join(intro_parts), 320)
        if intro_answer:
            faq_entries.append(
                {
                    "id": f"faq_spot_{counter:03d}",
                    "category": "景点介绍",
                    "aliases": [f"{spot_name}是什么", f"{spot_name}介绍", f"{spot_name}有什么亮点"],
                    "answer": intro_answer,
                    "sources": [base_source],
                }
            )
            counter += 1

        if record["具体位置"]:
            faq_entries.append(
                {
                    "id": f"faq_spot_{counter:03d}",
                    "category": "景点位置",
                    "aliases": [f"{spot_name}在哪里", f"{spot_name}位置"],
                    "answer": truncate(f"{spot_name}位于{record['具体位置']}", 240),
                    "sources": [base_source],
                }
            )
            counter += 1

        parameters = record["建筑/景观参数"]
        if parameters:
            aliases = [f"{spot_name}建筑参数", f"{spot_name}有什么参数"]
            if spot_name == "灵山大佛":
                aliases = ["灵山大佛有多高", "灵山大佛多高", "灵山大佛高度多少"] + aliases
            faq_entries.append(
                {
                    "id": f"faq_spot_{counter:03d}",
                    "category": "景点参数",
                    "aliases": aliases,
                    "answer": truncate(f"{spot_name}{parameters}", 260),
                    "sources": [base_source],
                }
            )
            counter += 1

        opening_info = record["演艺/开放信息"]
        if opening_info:
            faq_entries.append(
                {
                    "id": f"faq_spot_{counter:03d}",
                    "category": "开放信息",
                    "aliases": [f"{spot_name}开放时间", f"{spot_name}什么时候开放"],
                    "answer": truncate(f"{spot_name}{opening_info}", 260),
                    "sources": [base_source],
                }
            )
            counter += 1

    return faq_entries


def deduplicate_faq_entries(entries: list[dict]) -> list[dict]:
    deduped: list[dict] = []
    seen: set[tuple[tuple[str, ...], str]] = set()

    for entry in entries:
        aliases = tuple(sorted(set(entry["aliases"])))
        key = (aliases, entry["answer"])
        if key in seen:
            continue
        seen.add(key)
        deduped.append({**entry, "aliases": list(aliases)})

    return deduped


def build_behavior_summary(xlsx_path: Path) -> dict:
    df = pd.read_excel(xlsx_path)
    total_rows = int(len(df))
    unique_attractions = int(df["attraction_name"].nunique())
    date_min = str(pd.to_datetime(df["visit_date"]).min().date())
    date_max = str(pd.to_datetime(df["visit_date"]).max().date())

    attraction_type_distribution = df["attraction_type"].astype(str).value_counts().head(10).to_dict()
    top_attractions = df["attraction_name"].astype(str).value_counts().head(10).to_dict()
    avg_satisfaction_by_type = (
        df.groupby("attraction_type")["satisfaction"].mean().round(2).sort_values(ascending=False).to_dict()
    )
    avg_total_cost_by_type = (
        df.groupby("attraction_type")["total_cost"].mean().round(2).sort_values(ascending=False).to_dict()
    )
    avg_group_size = round(float(mean(df["group_size"].tolist())), 2)

    return {
        "dataset_name": xlsx_path.name,
        "intended_usage": "后台游客行为分析示例数据，不作为灵山胜境景区问答事实源。",
        "row_count": total_rows,
        "column_count": int(len(df.columns)),
        "date_range": {"start": date_min, "end": date_max},
        "unique_attraction_count": unique_attractions,
        "contains_lingshan_records": False,
        "attraction_type_distribution_top10": attraction_type_distribution,
        "top_attractions_by_visit_count": top_attractions,
        "average_satisfaction_by_type": avg_satisfaction_by_type,
        "average_total_cost_by_type": avg_total_cost_by_type,
        "average_group_size": avg_group_size,
        "columns": [str(column) for column in df.columns],
    }


def write_json(path: Path, payload) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def main() -> None:
    structured_docx = load_docx_by_keyword("结构化数据集")
    guide_docx = load_docx_by_keyword("游览指南")
    behavior_xlsx = load_xlsx()

    spot_records = extract_spot_records(structured_docx)
    guide_sections = extract_guide_sections(guide_docx)
    routes = extract_routes(guide_docx)
    guide_paragraphs = load_clean_paragraphs(guide_docx)
    guide_mentions = build_guide_mentions(guide_paragraphs, [record["景点名称"] for record in spot_records])

    knowledge_entries = [
        spot_record_to_knowledge_entry(record, structured_docx.name, guide_mentions)
        for record in spot_records
    ] + [
        guide_section_to_knowledge_entry(section, guide_docx.name)
        for section in guide_sections
        if section["content"] and section["title"] not in GENERIC_ROUTE_SECTION_TITLES
    ]

    faq_entries = deduplicate_faq_entries(
        build_overview_faq(guide_sections, guide_docx.name)
        + build_spot_faqs(spot_records, structured_docx.name, guide_mentions)
    )

    behavior_summary = build_behavior_summary(behavior_xlsx)

    write_json(PROCESSED_DIR / "knowledge_entries.json", knowledge_entries)
    write_json(PROCESSED_DIR / "guide_sections.json", guide_sections)
    write_json(PROCESSED_DIR / "route_recommendations.json", routes)
    write_json(PROCESSED_DIR / "faq_entries.json", faq_entries)
    write_json(PROCESSED_DIR / "visitor_behavior_summary.json", behavior_summary)

    print(
        json.dumps(
            {
                "knowledge_entries": len(knowledge_entries),
                "guide_sections": len(guide_sections),
                "routes": len(routes),
                "faq_entries": len(faq_entries),
                "behavior_rows": behavior_summary["row_count"],
            },
            ensure_ascii=False,
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
