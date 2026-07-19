"""将部署说明文档 Markdown 转换为规范的 .docx 文件。"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import sys

SRC = r"D:\桌面\软件杯\docs\交付物-部署说明文档.md"
DST = r"D:\桌面\软件杯\output\pdf\灵山胜境AI数字人导览系统-部署说明文档.docx"

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)

def set_no_border(cell):
    """移除单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_run(paragraph, text, bold=False, italic=False, size=None, color=None, font_name=None):
    """添加格式化的 run"""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color
    if font_name:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return run

def parse_inline(text, paragraph, base_size=None, base_color=None):
    """解析行内格式：**粗体**、*斜体*、`代码`、[链接]"""
    # 简单处理常见的 markdown 行内格式
    # 处理 **粗体**
    parts = []
    remaining = text
    while remaining:
        # 粗体 **text**
        bold_match = re.match(r'(.*?)\*\*(.+?)\*\*(.*)', remaining, re.DOTALL)
        if bold_match:
            if bold_match.group(1):
                parts.append(('normal', bold_match.group(1)))
            parts.append(('bold', bold_match.group(2)))
            remaining = bold_match.group(3)
            continue
        # 行内代码 `code`
        code_match = re.match(r'(.*?)`(.+?)`(.*)', remaining, re.DOTALL)
        if code_match:
            if code_match.group(1):
                parts.append(('normal', code_match.group(1)))
            parts.append(('code', code_match.group(2)))
            remaining = code_match.group(3)
            continue
        # 链接 [text](url)
        link_match = re.match(r'(.*?)\[(.+?)\]\((.+?)\)(.*)', remaining, re.DOTALL)
        if link_match:
            if link_match.group(1):
                parts.append(('normal', link_match.group(1)))
            parts.append(('normal', link_match.group(2)))
            remaining = link_match.group(4)
            continue
        parts.append(('normal', remaining))
        remaining = ''

    for kind, content in parts:
        if not content:
            continue
        if kind == 'bold':
            add_run(paragraph, content, bold=True, size=base_size, color=base_color)
        elif kind == 'code':
            r = add_run(paragraph, content, size=Pt(base_size.pt - 1 if base_size else 9), color=base_color)
            r.font.name = 'Consolas'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
        else:
            add_run(paragraph, content, size=base_size, color=base_color)

def add_heading_styled(doc, text, level):
    """添加标题并设置样式"""
    if level == 0:
        # 文档主标题
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(12)
        add_run(p, text, bold=True, size=Pt(22), color=RGBColor(0x1A, 0x1A, 0x1A))
    elif level == 1:
        h = doc.add_heading('', level=1)
        h.paragraph_format.space_before = Pt(24)
        h.paragraph_format.space_after = Pt(8)
        add_run(h, text, bold=True, size=Pt(18), color=RGBColor(0x2B, 0x57, 0x9A))
    elif level == 2:
        h = doc.add_heading('', level=2)
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        add_run(h, text, bold=True, size=Pt(15), color=RGBColor(0x33, 0x66, 0x99))
    elif level == 3:
        h = doc.add_heading('', level=3)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        add_run(h, text, bold=True, size=Pt(13), color=RGBColor(0x44, 0x77, 0xAA))

def parse_table_row(line):
    """解析表格行，返回单元格列表"""
    cells = line.strip().strip('|').split('|')
    return [c.strip() for c in cells]

def build_table(doc, rows):
    """根据表数据创建表格"""
    if len(rows) < 2:
        return
    # 跳过分隔行（如 |---|:---:|---|）
    data_rows = []
    for r in rows:
        cells = parse_table_row(r)
        # 跳过分隔行（所有单元格都由 -、:、空格组成）
        if all(re.match(r'^[-: ]+$', c) for c in cells):
            continue
        data_rows.append(cells)

    if not data_rows:
        return

    num_cols = len(data_rows[0])
    table = doc.add_table(rows=len(data_rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_cells in enumerate(data_rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_cells):
            if j >= num_cols:
                break
            cell = row.cells[j]
            # 清除默认空段落
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)

            if i == 0:
                # 表头行
                add_run(p, cell_text, bold=True, size=Pt(10), color=RGBColor(0xFF, 0xFF, 0xFF))
                set_cell_shading(cell, "2B579A")
            else:
                add_run(p, cell_text, size=Pt(10))
                if i % 2 == 0:
                    set_cell_shading(cell, "E8EFF7")

    doc.add_paragraph()  # 表格后空一行


def add_code_block(doc, lines):
    """添加代码块"""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(1)
        p.style = doc.styles['Normal']
        # 设置背景
        pPr = p._p.get_or_add_pPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F4F4F5')
        shading.set(qn('w:val'), 'clear')
        pPr.append(shading)
        add_run(p, line, size=Pt(9), font_name='Consolas', color=RGBColor(0x33, 0x33, 0x33))

    doc.add_paragraph()  # 代码块后空行


def add_blockquote(doc, lines):
    """添加引用块"""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(1)
        # 左边加竖线效果
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '12')
        left.set(qn('w:space'), '4')
        left.set(qn('w:color'), '2B579A')
        pBdr.append(left)
        pPr.append(pBdr)
        # 去掉 > 前缀
        text = re.sub(r'^>\s*', '', line)
        parse_inline(text, p, base_size=Pt(10), base_color=RGBColor(0x55, 0x55, 0x55))


def convert_md_to_docx(md_path, docx_path):
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(10.5)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # 跳过文件开头的 frontmatter
        if line == '---' and i < 3:
            i += 1
            while i < len(lines) and lines[i].strip() != '---':
                i += 1
            i += 1
            continue

        # 水平分隔线
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # 标题
        heading_match = re.match(r'^(#{1,3})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            add_heading_styled(doc, text, level)
            i += 1
            continue

        # 代码块开始 ```
        if line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines):
                if lines[i].strip().startswith('```'):
                    i += 1
                    break
                code_lines.append(lines[i].rstrip())
                i += 1
            add_code_block(doc, code_lines)
            continue

        # 引用块 >
        if line.strip().startswith('>'):
            quote_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('>'):
                quote_lines.append(lines[i])
                i += 1
            add_blockquote(doc, quote_lines)
            continue

        # 表格（检测 | 开头的行）
        if line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            build_table(doc, table_lines)
            continue

        # 无序列表
        list_match = re.match(r'^(\s*)-\s+(.+)$', line)
        if list_match:
            indent_level = len(list_match.group(1)) // 2
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.5 + indent_level * 0.5)
            p.paragraph_format.first_line_indent = Cm(-0.4)
            add_run(p, '• ', bold=True, size=Pt(10.5))
            parse_inline(list_match.group(2), p, base_size=Pt(10.5))
            i += 1
            continue

        # 有序列表
        ordered_match = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
        if ordered_match:
            indent_level = len(ordered_match.group(1)) // 2
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.5 + indent_level * 0.5)
            p.paragraph_format.first_line_indent = Cm(-0.4)
            parse_inline(ordered_match.group(2), p, base_size=Pt(10.5))
            i += 1
            continue

        # checkbox 列表 - [ ] 或 - [x]
        checkbox_match = re.match(r'^(\s*)-\s+\[([ xX])\]\s+(.+)$', line)
        if checkbox_match:
            checked = checkbox_match.group(2) in ('x', 'X')
            symbol = '☑' if checked else '☐'
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.5)
            add_run(p, f'{symbol} ', bold=True, size=Pt(10.5),
                    color=RGBColor(0x2B, 0x57, 0x9A) if checked else None)
            parse_inline(checkbox_match.group(3), p, base_size=Pt(10.5))
            i += 1
            continue

        # 空行
        if not line.strip():
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        parse_inline(line, p, base_size=Pt(10.5))
        i += 1

    # 保存
    doc.save(docx_path)
    print(f"已生成: {docx_path}")


if __name__ == '__main__':
    convert_md_to_docx(SRC, DST)
