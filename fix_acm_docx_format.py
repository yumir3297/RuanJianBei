from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


SOURCE = Path(r"D:\xwechat_files\wxid_uft3lonxjh1g22_4c8b\msg\file\2026-06\Luo_Liu_Liu_ICISS2026_ACM_template(4).docx")
OUTPUT = Path(r"D:\桌面\软件杯\Luo_Liu_Liu_ICISS2026_ACM_template(4)_formatfix.docx")

INSTITUTION = "School of Engineering, Nanfang College Guangzhou, Guangzhou 510971, China"
AUTHOR_BLOCK = [
    ("Authors", "Yuewen Luo"),
    ("Affiliation", f"{INSTITUTION}, liyumir087@gmail.com"),
    ("Authors", "Kairui Liu"),
    ("Affiliation", f"{INSTITUTION}, 250162liukr@stu.nfu.edu.cn"),
    ("Authors", "Chenglian Liu*"),
    ("Affiliation", f"{INSTITUTION}, chenglian.liu@gmail.com"),
    ("Affiliation", "*Corresponding author: Chenglian Liu."),
]

NOTE_LINES = [
    "Notes:",
    "*This form helps the conference understand the paper better.",
    "*All author information should be updated by the authors before final submission if a title or personal website must be supplied.",
]


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.text = text
    return new_para


def iter_all_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def clear_run_formatting(doc: Document) -> None:
    for paragraph in iter_all_paragraphs(doc):
        for run in paragraph.runs:
            run.bold = None
            run.italic = None
            run.underline = None
            run.font.name = None
            run.font.size = None
            try:
                run.font.color.rgb = None
            except Exception:
                pass


def clear_paragraph_formatting(doc: Document) -> None:
    for paragraph in iter_all_paragraphs(doc):
        fmt = paragraph.paragraph_format
        fmt.left_indent = None
        fmt.right_indent = None
        fmt.first_line_indent = None
        fmt.space_before = None
        fmt.space_after = None
        fmt.line_spacing = None


def is_formula_table(table: Table) -> bool:
    if len(table.rows) != 1 or len(table.columns) != 2:
        return False

    right_text = "".join(p.text for p in table.rows[0].cells[1].paragraphs).strip()
    if not re.fullmatch(r"\(\d+\)", right_text):
        return False

    left_tc = table.rows[0].cells[0]._tc
    has_math = any(el.tag in {qn("m:oMath"), qn("m:oMathPara")} for el in left_tc.iter())
    return has_math


def build_display_formula_paragraph(doc: Document, table: Table) -> OxmlElement:
    style_id = doc.styles["DisplayFormula"].style_id
    new_p = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    p_style = OxmlElement("w:pStyle")
    p_style.set(qn("w:val"), style_id)
    p_pr.append(p_style)
    new_p.append(p_pr)

    left_cell = table.rows[0].cells[0]
    for cell_paragraph in left_cell.paragraphs:
        for child in cell_paragraph._p:
            if child.tag == qn("w:pPr"):
                continue
            if child.tag == qn("m:oMathPara"):
                for math_child in child:
                    if math_child.tag == qn("m:oMath"):
                        new_p.append(deepcopy(math_child))
            elif child.tag == qn("m:oMath"):
                new_p.append(deepcopy(child))
            elif child.tag == qn("w:r"):
                text = "".join(t.text for t in child.iter(qn("w:t")) if t.text).strip()
                if text:
                    new_p.append(deepcopy(child))

    number_text = "".join(p.text for p in table.rows[0].cells[1].paragraphs).strip()
    if number_text:
        run = OxmlElement("w:r")
        run.append(OxmlElement("w:tab"))
        text = OxmlElement("w:t")
        text.text = number_text
        run.append(text)
        new_p.append(run)

    return new_p


def convert_formula_tables(doc: Document) -> tuple[int, int]:
    converted = 0
    removed_empty = 0
    body = doc._body._element

    for child in list(body.iterchildren()):
        if child.tag != qn("w:tbl"):
            continue

        table = Table(child, doc._body)
        if not is_formula_table(table):
            continue

        display_formula = build_display_formula_paragraph(doc, table)
        child.addprevious(display_formula)

        next_el = child.getnext()
        while next_el is not None and next_el.tag == qn("w:p"):
            paragraph = Paragraph(next_el, doc._body)
            if paragraph.text.strip():
                break
            to_remove = next_el
            next_el = next_el.getnext()
            body.remove(to_remove)
            removed_empty += 1

        body.remove(child)
        converted += 1

    return converted, removed_empty


def rebuild_author_block(doc: Document) -> None:
    title_para, short_title_para = doc.paragraphs[0], doc.paragraphs[1]
    _ = title_para, short_title_para

    first_author = doc.paragraphs[2]
    first_affiliation = doc.paragraphs[3]
    second_author = doc.paragraphs[4]

    first_author.style = AUTHOR_BLOCK[0][0]
    first_author.text = AUTHOR_BLOCK[0][1]

    first_affiliation.style = AUTHOR_BLOCK[1][0]
    first_affiliation.text = AUTHOR_BLOCK[1][1]

    second_author.style = AUTHOR_BLOCK[2][0]
    second_author.text = AUTHOR_BLOCK[2][1]

    current = second_author
    for style, text in AUTHOR_BLOCK[3:]:
        current = insert_paragraph_after(current, text=text, style=style)


def rebuild_background_notes(doc: Document) -> None:
    background_title = next(p for p in doc.paragraphs if p.text.strip() == "Authors' Background")
    background_title.style = "Paper-Title"

    note_paragraph = next(p for p in doc.paragraphs if p.text.strip().startswith("Notes:"))
    note_paragraph.style = "Paper-Title"
    note_paragraph.text = NOTE_LINES[0]

    current = note_paragraph
    for line in NOTE_LINES[1:]:
        current = insert_paragraph_after(current, text=line, style="Paper-Title")


def main() -> None:
    doc = Document(str(SOURCE))

    rebuild_author_block(doc)
    rebuild_background_notes(doc)
    converted, removed_empty = convert_formula_tables(doc)
    clear_run_formatting(doc)
    clear_paragraph_formatting(doc)

    doc.save(str(OUTPUT))
    print(f"saved={OUTPUT}")
    print(f"formula_tables_converted={converted}")
    print(f"empty_paragraphs_removed={removed_empty}")


if __name__ == "__main__":
    main()
